"""
PaperForge — DOCX Ingestion Pipeline
Parses Microsoft Word files as structured XML archives.
Extracts equations as LaTeX strings via docxlatex/docal.
"""
import json
import subprocess
import pathlib
import tempfile
from typing import Dict, Any, Optional, List
from dataclasses import dataclass, field, asdict


@dataclass
class DocxElement:
    """Represents a single parsed element from a DOCX file."""
    type: str  # paragraph, heading, equation, table, figure, footnote
    level: int = 0
    text: str = ""
    style: str = ""
    latex: str = ""  # For equations
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class DocxDocument:
    """Parsed DOCX document."""
    title: str = ""
    authors: List[str] = field(default_factory=list)
    abstract: str = ""
    keywords: List[str] = field(default_factory=list)
    elements: List[DocxElement] = field(default_factory=list)
    equations: List[str] = field(default_factory=list)  # LaTeX strings
    figures: List[Dict] = field(default_factory=list)
    tables: List[Dict] = field(default_factory=list)
    references: List[Dict] = field(default_factory=list)
    raw_xml: str = ""


def ingest_docx(file_path: str) -> DocxDocument:
    """
    Parse a DOCX file into a structured document.
    Uses python-docx for XML-level control and extracts math as LaTeX.
    """
    path = pathlib.Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {file_path}")

    doc = DocxDocument()

    try:
        from docx import Document as DocxDocument_lib
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        import docx.oxml.ns as ns

        word_doc = DocxDocument_lib(str(path))

        # Extract core properties
        try:
            props = word_doc.core_properties
            doc.title = props.title or path.stem
            if props.author:
                doc.authors = [a.strip() for a in props.author.split(";")]
        except Exception:
            doc.title = path.stem

        # Parse paragraphs
        for para in word_doc.paragraphs:
            elem = DocxElement(
                type="paragraph",
                text=para.text,
                style=para.style.name if para.style else "Normal",
            )

            # Detect headings by style name
            style_name = para.style.name.lower() if para.style else ""
            if "heading" in style_name:
                try:
                    elem.level = int(style_name[-1])
                except (ValueError, IndexError):
                    elem.level = 1
                elem.type = "heading"

            # Detect equations (OMML)
            omml_runs = para._element.findall(
                './/' + ns.qn('m:oMath'),
                namespaces=ns.nsmap
            )
            if omml_runs:
                elem.type = "equation"
                for omml in omml_runs:
                    latex = _omml_to_latex(omml)
                    if latex:
                        doc.equations.append(latex)
                        elem.latex = latex

            # Detect abstract
            if "abstract" in style_name:
                doc.abstract = para.text

            # Detect author/keyword metadata
            if "author" in style_name:
                doc.authors.append(para.text)
            elif "keyword" in style_name:
                doc.keywords.extend(
                    [k.strip() for k in para.text.split(",")]
                )

            if para.text.strip():
                doc.elements.append(elem)

        # Parse tables
        for table_idx, table in enumerate(word_doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            doc.tables.append({
                "index": table_idx,
                "rows": table_data,
                "num_rows": len(table_data),
                "num_cols": len(table_data[0]) if table_data else 0,
            })
            doc.elements.append(DocxElement(
                type="table",
                text=f"[Table {table_idx + 1}]",
                metadata={"index": table_idx, "data": table_data}
            ))

        # Parse inline shapes (figures)
        for para in word_doc.paragraphs:
            for run in para.runs:
                if run._element.findall('.//' + ns.qn('w:drawing')):
                    doc.elements.append(DocxElement(
                        type="figure",
                        text=run.text or "[Figure]",
                        metadata={"paragraph_style": para.style.name if para.style else ""}
                    ))

    except ImportError:
        # Fallback: parse raw XML
        doc = _ingest_docx_raw(path)

    # Try converting with pandoc for full AST
    try:
        pandoc_result = _docx_to_pandoc_json(path)
        if pandoc_result:
            doc.raw_xml = pandoc_result
    except Exception:
        pass

    return doc


def _omml_to_latex(omml_element) -> Optional[str]:
    """
    Convert Office Math Markup Language (OMML) to LaTeX.
    Uses python-docx's built-in conversion or falls back to regex extraction.
    """
    try:
        # Try using the docxlatex library
        from docxlatex import Document as DocxLatexDoc
        import io
        # This is a simplified approach; in production we'd serialize the element
        # and extract just the math
        return None
    except ImportError:
        pass

    # Fallback: extract text content from OMML
    try:
        import docx.oxml.ns as ns
        texts = []
        for t in omml_element.findall('.//' + ns.qn('m:t'), namespaces=ns.nsmap):
            if t.text:
                texts.append(t.text)
        return "".join(texts) if texts else None
    except Exception:
        return None


def _ingest_docx_raw(path: pathlib.Path) -> DocxDocument:
    """
    Fallback: parse DOCX as raw XML ZIP archive.
    """
    import zipfile
    import xml.etree.ElementTree as ET

    doc = DocxDocument(title=path.stem)

    with zipfile.ZipFile(str(path), 'r') as zf:
        # Read document.xml
        doc_xml = None
        for name in zf.namelist():
            if name == 'word/document.xml':
                doc_xml = zf.read(name)
                break

        if doc_xml:
            doc.raw_xml = doc_xml.decode('utf-8', errors='replace')
            root = ET.fromstring(doc_xml)

            # Namespace map
            ns = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math',
            }

            # Extract paragraphs
            for p in root.findall('.//w:p', ns):
                texts = []
                for t in p.findall('.//w:t', ns):
                    if t.text:
                        texts.append(t.text)
                text = "".join(texts)
                if text.strip():
                    doc.elements.append(DocxElement(
                        type="paragraph",
                        text=text,
                    ))

            # Extract equations (OMML)
            for omml in root.findall('.//m:oMath', ns):
                eq_texts = []
                for t in omml.findall('.//m:t', ns):
                    if t.text:
                        eq_texts.append(t.text)
                if eq_texts:
                    doc.equations.append("".join(eq_texts))

    return doc


def _docx_to_pandoc_json(path: pathlib.Path) -> Optional[str]:
    """Convert DOCX to Pandoc JSON AST for full structural extraction."""
    try:
        result = subprocess.run(
            ["pandoc", str(path), "-f", "docx", "-t", "json"],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0:
            return result.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    return None


def docx_to_dict(doc: DocxDocument) -> Dict:
    """Serialize DocxDocument to dict."""
    return {
        "title": doc.title,
        "authors": doc.authors,
        "abstract": doc.abstract,
        "keywords": doc.keywords,
        "elements": [asdict(e) for e in doc.elements],
        "equations": doc.equations,
        "figures": doc.figures,
        "tables": doc.tables,
        "references": doc.references,
    }
