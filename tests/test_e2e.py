#!/usr/bin/env python3
"""
PaperForge — End-to-End Pipeline Tests
Tests the full pipeline with real document generation and conversion.
"""
import sys
import json
import pathlib
import tempfile

PROJECT_ROOT = pathlib.Path(__file__).parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

PASS = "✅ PASS"
FAIL = "❌ FAIL"


def test(name, passed, msg=""):
    status = PASS if passed else FAIL
    print(f"  {status} {name}" + (f" — {msg}" if msg else ""))
    return passed


def create_test_docx(path: str):
    """Create a test DOCX file with equations, tables, and figures."""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading('PaperForge End-to-End Test Paper', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Authors
    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = authors.add_run('Alice Researcher, Bob Scientist')
    run.font.size = Pt(12)

    # Abstract
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        'This is a test paper for the PaperForge document conversion pipeline. '
        'It contains equations, tables, and structured content to verify '
        'the ingestion, transformation, and export phases.'
    )

    # Keywords
    kw = doc.add_paragraph()
    kw.add_run('Keywords: ').bold = True
    kw.add_run('testing, pipeline, document conversion')

    # Introduction
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        'The PaperForge system converts academic papers between formats. '
        'Consider the fundamental equation E = mc², proposed by Einstein.'
    )

    # Methods with equation
    doc.add_heading('Methods', level=1)
    doc.add_paragraph(
        'Our approach uses the following optimization: '
    )
    # Add equation paragraph
    eq_para = doc.add_paragraph()
    eq_run = eq_para.add_run('minimize f(x) = ∫₀^∞ exp(-x²) dx')
    eq_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Results with table
    doc.add_heading('Results', level=1)
    doc.add_paragraph('Table 1 shows our experimental results.')

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ['Method', 'Accuracy', 'Time (ms)']
    data = [
        ['Baseline', '85.2%', '120'],
        ['PaperForge', '97.8%', '45'],
        ['Proposed', '98.5%', '38'],
    ]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r, row_data in enumerate(data):
        for c, val in enumerate(row_data):
            table.rows[r + 1].cells[c].text = val

    # Discussion
    doc.add_heading('Discussion', level=1)
    doc.add_paragraph(
        'The results demonstrate that our approach outperforms existing methods. '
        'The Fourier transform is defined as: '
    )
    eq2 = doc.add_paragraph()
    eq2.add_run('f̂(ξ) = ∫_{-∞}^{∞} f(x) e^{-2πixξ} dx')
    eq2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Conclusion
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'PaperForge provides a robust pipeline for academic paper conversion. '
        'Future work includes AI-assisted editing and real-time collaboration.'
    )

    doc.save(path)


def main():
    all_passed = True
    print("=" * 60)
    print("  PaperForge — End-to-End Pipeline Test Suite")
    print("=" * 60)

    with tempfile.TemporaryDirectory() as tmpdir:
        tmp = pathlib.Path(tmpdir)

        # ────────────────────────────────────────────────────────
        # Test 1: DOCX Creation & Ingestion
        # ────────────────────────────────────────────────────────
        print("\n[1] DOCX Ingestion")
        try:
            docx_path = str(tmp / "test_paper.docx")
            create_test_docx(docx_path)

            from src.ingestion.pipeline import ingest, detect_format

            fmt = detect_format(docx_path)
            test("Format detection", fmt == "docx", f"Detected: {fmt}")

            doc = ingest(docx_path, "docx")
            test("Title extraction", "PaperForge" in doc.title, f"Title: {doc.title}")
            test("Author extraction", len(doc.authors) >= 1, f"Authors: {doc.authors}")
            test("Elements parsed", len(doc.elements) > 5, f"Elements: {len(doc.elements)}")
            test("Tables detected", len(doc.tables) >= 1, f"Tables: {len(doc.tables)}")

        except Exception as e:
            test("DOCX ingestion", False, str(e))
            all_passed = False

        # ────────────────────────────────────────────────────────
        # Test 2: Markdown Ingestion
        # ────────────────────────────────────────────────────────
        print("\n[2] Markdown Ingestion")
        try:
            md_path = str(tmp / "test_paper.md")
            pathlib.Path(md_path).write_text("""---
title: "Markdown Test Paper"
author: "Charlie Test"
abstract: "Testing markdown ingestion."
keywords: "test, markdown"
---

# Introduction

This paper tests markdown ingestion with equations.

## Methods

The energy equation is $E = mc^2$.

The Gaussian integral:

$$\\int_{-\\infty}^{\\infty} e^{-x^2} dx = \\sqrt{\\pi}$$

## Results

| Method | Score |
|--------|-------|
| A | 95.2 |
| B | 97.8 |

![Test Figure](figure.png)

# Conclusion

Markdown ingestion works correctly.
""")

            doc = ingest(md_path, "markdown")
            test("Title extraction", "Markdown Test Paper" in doc.title)
            test("Author extraction", len(doc.authors) >= 1)
            test("Equation extraction", len(doc.equations) >= 2, f"Equations: {len(doc.equations)}")
            test("Table detection", len(doc.tables) >= 1)
            test("Figure detection", len(doc.figures) >= 1, f"Figures: {len(doc.figures)}")

        except Exception as e:
            test("Markdown ingestion", False, str(e))
            all_passed = False

        # ────────────────────────────────────────────────────────
        # Test 3: LaTeX Ingestion
        # ────────────────────────────────────────────────────────
        print("\n[3] LaTeX Ingestion")
        try:
            tex_path = str(tmp / "test_paper.tex")
            pathlib.Path(tex_path).write_text(r"""\documentclass{article}
\usepackage{amsmath}
\title{LaTeX Test Paper}
\author{Diana Researcher}
\begin{document}
\begin{abstract}
Testing LaTeX ingestion with equations.
\end{abstract}
\keywords{test, latex}

\section{Introduction}
Consider Maxwell's equations.

\section{Equations}
\begin{equation}
\nabla \cdot \mathbf{E} = \frac{\rho}{\epsilon_0}
\end{equation}

\begin{equation}
\nabla \times \mathbf{B} = \mu_0 \mathbf{J} + \mu_0 \epsilon_0 \frac{\partial \mathbf{E}}{\partial t}
\end{equation}

\section{Conclusion}
LaTeX parsing works correctly.

\bibliography{refs}
\end{document}
""")

            doc = ingest(tex_path, "latex")
            test("Title extraction", "LaTeX Test Paper" in doc.title)
            test("Author extraction", len(doc.authors) >= 1)
            test("Section extraction", len(doc.sections) >= 3, f"Sections: {len(doc.sections)}")
            test("Equation extraction", len(doc.equations) >= 2, f"Equations: {len(doc.equations)}")
            test("Pandoc AST available", len(doc.pandoc_json) > 100)

        except Exception as e:
            test("LaTeX ingestion", False, str(e))
            all_passed = False

        # ────────────────────────────────────────────────────────
        # Test 4: Transformation Engine
        # ────────────────────────────────────────────────────────
        print("\n[4] Transformation Engine")
        try:
            from src.transformation.engine import (
                transform, load_style_map, normalize_math,
                get_default_style_map, apply_lua_filters
            )

            # Build a test AST (Pandoc API format)
            ast = {
                "pandoc-api-version": [1, 23, 1],
                "meta": {
                    "title": {"t": "MetaInlines", "c": [{"t": "Str", "c": "Test"}]},
                    "author": {"t": "MetaList", "c": [
                        {"t": "MetaInlines", "c": [{"t": "Str", "c": "Author"}]}
                    ]},
                },
                "blocks": [
                    {"t": "Header", "c": [1, ["intro", [], []], [{"t": "Str", "c": "Intro"}]]},
                    {"t": "Para", "c": [{"t": "Str", "c": "Text with "}, {
                        "t": "Math",
                        "c": [{"t": "InlineMath"}, "E = mc^2"]
                    }]},
                    {"t": "Para", "c": [{"t": "Math", "c": [{"t": "DisplayMath"}, "\\int_0^\\infty e^{-x} dx = 1"]}]},
                ]
            }
            ast_json = json.dumps(ast)

            # Test style map loading
            for journal in ["ieee", "springer", "wiley", "elsevier", "nature", "acm", "mdpi"]:
                sm = get_default_style_map(journal)
                test(f"Style map: {journal}", "class" in sm and "font" in sm)

            # Test math normalization
            norm = normalize_math(r"\begin{align} x &= 1 \\ y &= 2 \end{align}")
            test("Math normalization", r'\begin{aligned}' in norm)

            # Test full transform
            result = transform(ast_json, "ieee")
            test("Transform completes", len(result.ast_json) > 0)
            test("Equations normalized", len(result.normalized_equations) >= 1)
            test("Style map loaded", len(result.style_map) > 0)

        except Exception as e:
            test("Transformation", False, str(e))
            all_passed = False

        # ────────────────────────────────────────────────────────
        # Test 5: Semantic Audit
        # ────────────────────────────────────────────────────────
        print("\n[5] Semantic Audit")
        try:
            from src.auditing.semantic import full_audit, audit_report

            audits = full_audit(ast_json)
            report = audit_report(audits)

            test("Audit completes", True)
            test("Metadata audit", "metadata" in report["audits"])
            test("Structure audit", "structure" in report["audits"])
            test("Math audit", "math" in report["audits"])
            test("Score calculated", report["score"] >= 0)

        except Exception as e:
            test("Semantic audit", False, str(e))
            all_passed = False

        # ────────────────────────────────────────────────────────
        # Test 6: Export Pipeline
        # ────────────────────────────────────────────────────────
        print("\n[6] Export Pipeline")
        try:
            from src.export.exporter import export_document

            # Export to DOCX
            docx_out = str(tmp / "output.docx")
            result = export_document(ast_json, docx_out, "docx", journal="ieee")
            test("DOCX export", result.success, result.error if not result.success else f"{result.size_bytes} bytes")

            # Export to LaTeX
            tex_out = str(tmp / "output.tex")
            result = export_document(ast_json, tex_out, "latex", journal="ieee")
            test("LaTeX export", result.success, result.error if not result.success else f"{result.size_bytes} bytes")

            # Export to HTML
            html_out = str(tmp / "output.html")
            result = export_document(ast_json, html_out, "html", journal="springer")
            test("HTML export", result.success, result.error if not result.success else f"{result.size_bytes} bytes")

        except Exception as e:
            test("Export", False, str(e))
            all_passed = False

        # ────────────────────────────────────────────────────────
        # Test 7: Full Pipeline (DOCX → IEEE DOCX)
        # ────────────────────────────────────────────────────────
        print("\n[7] Full Pipeline: DOCX → IEEE DOCX")
        try:
            from pipeline.translator import run_pipeline

            output_dir = str(tmp / "pipeline_output")
            result = run_pipeline(
                input_file=docx_path,
                journal="ieee",
                output_format="docx",
                output_dir=output_dir,
                quiet=True,
            )

            test("Pipeline completes", result["success"])
            test("All steps done", all(
                s == "completed" for s in result["steps"].values()
            ), f"Steps: {result['steps']}")
            test("Output exists", len(result.get("outputs", {})) > 0,
                 f"Outputs: {list(result.get('outputs', {}).keys())}")

        except Exception as e:
            test("Full pipeline", False, str(e))
            all_passed = False

        # ────────────────────────────────────────────────────────
        # Test 8: Full Pipeline (Markdown → Springer HTML)
        # ────────────────────────────────────────────────────────
        print("\n[8] Full Pipeline: MD → Springer HTML")
        try:
            output_dir = str(tmp / "pipeline_md_output")
            result = run_pipeline(
                input_file=md_path,
                journal="springer",
                output_format="html",
                output_dir=output_dir,
                quiet=True,
            )

            test("Pipeline completes", result["success"])
            test("HTML output", "html" in result.get("outputs", {}))

        except Exception as e:
            test("MD→HTML pipeline", False, str(e))
            all_passed = False

    # Summary
    print(f"\n{'=' * 60}")
    if all_passed:
        print(f"  ✅ ALL E2E TESTS PASSED")
    else:
        print(f"  ❌ SOME E2E TESTS FAILED")
    print(f"{'=' * 60}\n")

    return 0 if all_passed else 1


if __name__ == "__main__":
    sys.exit(main())
